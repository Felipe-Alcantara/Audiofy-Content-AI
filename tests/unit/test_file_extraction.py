"""Testes da extração local de texto de arquivos (sem custo de IA)."""

import sys
import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

sys.path.insert(0, str(Path(__file__).resolve().parents[2] / "src"))

from audiofy.file_extraction import (  # noqa: E402
    MAX_FILE_BYTES,
    ExtractionResult,
    _epub_html_to_text,
    _extract_image,
    _normalize,
    _ocr_languages,
    _ocr_pdf,
    _strip_running_headers,
    _title_from_path,
    extract_file,
    ocr_available,
)

LONGO = (
    "Primeiro parágrafo com conteúdo suficientemente longo para ultrapassar "
    "o limite mínimo exigido pela extração automática de texto.\n\n"
    "Segundo parágrafo, também extenso, para garantir que a normalização "
    "preserve as quebras entre blocos."
)


class NormalizacaoTest(unittest.TestCase):
    def test_converte_quebras_windows_e_colapsa_linhas_extras(self):
        self.assertEqual(_normalize("a\r\n\r\n\r\n\r\nb"), "a\n\nb")

    def test_remove_espacos_antes_da_quebra_e_nas_bordas(self):
        self.assertEqual(_normalize("  linha   \n  outra  "), "linha\n  outra")

    def test_titulo_derivado_do_nome_troca_separadores_por_espaco(self):
        self.assertEqual(_title_from_path(Path("/tmp/meu_livro-final.pdf")), "meu livro final")


class TextoSimplesTest(unittest.TestCase):
    def setUp(self):
        self.directory = tempfile.TemporaryDirectory()
        self.addCleanup(self.directory.cleanup)
        self.root = Path(self.directory.name)

    def test_extrai_txt_utf8(self):
        path = self.root / "conto.txt"
        path.write_text(LONGO, encoding="utf-8")
        result = extract_file(path)
        self.assertEqual(result.method, "plain-text")
        self.assertFalse(result.needs_fallback)
        self.assertEqual(result.title, "conto")
        self.assertIn("Primeiro parágrafo", result.text)

    def test_extrai_markdown_com_acentos_em_latin1(self):
        path = self.root / "notas.md"
        path.write_bytes(LONGO.encode("latin-1"))
        result = extract_file(path)
        self.assertFalse(result.needs_fallback)
        self.assertIn("conteúdo", result.text)

    def test_arquivo_vazio_e_recusado(self):
        path = self.root / "vazio.txt"
        path.write_text("", encoding="utf-8")
        with self.assertRaisesRegex(ValueError, "vazio"):
            extract_file(path)

    def test_arquivo_inexistente_falha(self):
        with self.assertRaises(FileNotFoundError):
            extract_file(self.root / "nao-existe.txt")

    def test_formato_desconhecido_falha_com_lista_de_aceitos(self):
        path = self.root / "planilha.xlsx"
        path.write_bytes(b"conteudo binario qualquer")
        with self.assertRaisesRegex(ValueError, "não suportado"):
            extract_file(path)

    def test_arquivo_gigante_e_recusado_antes_de_ler(self):
        path = self.root / "livro.pdf"
        path.write_bytes(b"x")
        real_stat = Path.stat

        def stat_gigante(self, **kwargs):
            info = real_stat(self, **kwargs)
            if self == path:
                return type(info)(
                    (info.st_mode, *tuple(info)[1:6], MAX_FILE_BYTES + 1, *tuple(info)[7:])
                )
            return info

        with patch.object(Path, "stat", stat_gigante):
            with self.assertRaisesRegex(ValueError, "200 MiB"):
                extract_file(path)


class DocxTest(unittest.TestCase):
    def setUp(self):
        self.directory = tempfile.TemporaryDirectory()
        self.addCleanup(self.directory.cleanup)
        self.root = Path(self.directory.name)

    def _build(self, name: str, paragraphs: list[str]) -> Path:
        from docx import Document

        document = Document()
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        path = self.root / name
        document.save(str(path))
        return path

    def test_extrai_paragrafos_do_docx(self):
        path = self._build("relatorio.docx", LONGO.split("\n\n"))
        result = extract_file(path)
        self.assertEqual(result.method, "python-docx")
        self.assertFalse(result.needs_fallback)
        self.assertIn("Primeiro parágrafo", result.text)

    def test_docx_sem_texto_pede_outra_via(self):
        path = self._build("so-imagens.docx", ["", "   "])
        result = extract_file(path)
        self.assertTrue(result.needs_fallback)
        self.assertEqual(result.text, "")
        self.assertIn("extraível", result.reason)

    def test_incorpora_o_texto_das_tabelas(self):
        from docx import Document

        document = Document()
        document.add_paragraph(LONGO)
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Receita"
        table.rows[0].cells[1].text = "1200"
        path = self.root / "com-tabela.docx"
        document.save(str(path))

        result = extract_file(path)
        self.assertIn("Receita — 1200", result.text)

    def test_docx_corrompido_vira_erro_legivel(self):
        path = self.root / "quebrado.docx"
        path.write_bytes(b"isto nao e um zip de docx")
        with self.assertRaisesRegex(ValueError, "abrir o DOCX"):
            extract_file(path)

    def test_usa_o_titulo_das_propriedades_quando_definido(self):
        from docx import Document

        document = Document()
        document.core_properties.title = "Título Oficial"
        document.add_paragraph(LONGO)
        path = self.root / "com-titulo.docx"
        document.save(str(path))

        self.assertEqual(extract_file(path).title, "Título Oficial")


class EpubTest(unittest.TestCase):
    def test_converte_html_do_capitulo_em_texto(self):
        markup = (
            "<html><body><h1>Capítulo</h1>"
            "<script>alert('fora')</script>"
            "<p>Primeiro&nbsp;bloco.</p><p>Segundo bloco.</p></body></html>"
        )
        text = _epub_html_to_text(markup)
        self.assertNotIn("alert", text)
        self.assertIn("Primeiro\xa0bloco.", text)
        self.assertIn("Segundo bloco.", text)

    def test_extrai_epub_real(self):
        from ebooklib import epub

        directory = tempfile.TemporaryDirectory()
        self.addCleanup(directory.cleanup)
        book = epub.EpubBook()
        book.set_identifier("id-teste")
        book.set_title("Meu Livro")
        book.set_language("pt")
        chapter = epub.EpubHtml(title="Cap 1", file_name="c1.xhtml", lang="pt")
        chapter.content = f"<html><body><p>{LONGO}</p></body></html>"
        book.add_item(chapter)
        book.spine = ["nav", chapter]
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        path = Path(directory.name) / "livro.epub"
        epub.write_epub(str(path), book)

        result = extract_file(path)
        self.assertEqual(result.method, "ebooklib")
        self.assertFalse(result.needs_fallback)
        self.assertEqual(result.title, "Meu Livro")
        self.assertIn("Primeiro parágrafo", result.text)


class RodapeRepetidoTest(unittest.TestCase):
    """Cabeçalhos e rodapés de diagramação viram ruído lido em voz alta.

    Caso real: um PDF de livro repetia o rodapé do InDesign em 23 páginas; uma
    delas só tinha isso, e o TTS não devolveu áudio para o trecho, derrubando a
    geração inteira depois de 15 falas já pagas.
    """

    def _paginas(self, corpos):
        # Páginas de livro têm várias linhas de miolo; o rodapé fecha cada uma.
        return [
            corpo + f"\n14909-Homenagem à Catalunha (4P).indd   {numero} 15/02/21   15:07"
            for numero, corpo in enumerate(corpos, 1)
        ]

    def _miolo(self, numero):
        return "\n".join(
            [
                f"Parágrafo de abertura da página {numero}, com texto corrido.",
                "Segunda linha do miolo, também com conteúdo real do livro.",
                "Terceira linha, para a página ter corpo além das bordas.",
                f"Fechamento do trecho {numero} antes do rodapé.",
            ]
        )

    def test_remove_rodape_que_so_muda_o_numero_da_pagina(self):
        paginas = self._paginas([self._miolo(n) for n in range(1, 9)])
        limpas = _strip_running_headers(paginas)
        self.assertEqual(sum(p.count(".indd") for p in limpas), 0)
        for numero in range(1, 9):
            self.assertIn(f"Parágrafo de abertura da página {numero},", "\n".join(limpas))

    def test_pagina_que_era_so_rodape_fica_vazia(self):
        # O caso que derrubou a geração real: a página 8 do PDF só tinha o rodapé.
        paginas = self._paginas([self._miolo(n) for n in range(1, 8)] + [""])
        limpas = _strip_running_headers(paginas)
        self.assertEqual(limpas[-1].strip(), "")

    def test_preserva_documentos_curtos_sem_analisar(self):
        paginas = ["Primeira página.", "Segunda página."]
        self.assertEqual(_strip_running_headers(paginas), paginas)

    def test_nao_remove_linha_repetida_no_miolo_da_pagina(self):
        # Refrão, epígrafe ou estribilho no meio do texto é conteúdo, não rodapé.
        paginas = [
            f"Abertura {n}\nlinha dois\nrefrão que se repete\npenúltima {n}\nFim {n}"
            for n in range(1, 9)
        ]
        limpas = _strip_running_headers(paginas)
        self.assertEqual(sum(p.count("refrão que se repete") for p in limpas), 8)

    def test_remove_cabecalho_alem_do_rodape(self):
        paginas = [f"TÍTULO DO LIVRO\nCorpo {n} da página.\npágina {n}" for n in range(1, 9)]
        limpas = _strip_running_headers(paginas)
        self.assertEqual(sum(p.count("TÍTULO DO LIVRO") for p in limpas), 0)
        self.assertIn("Corpo 4 da página.", "\n".join(limpas))


class PdfTest(unittest.TestCase):
    def setUp(self):
        self.directory = tempfile.TemporaryDirectory()
        self.addCleanup(self.directory.cleanup)
        self.root = Path(self.directory.name)

    def _empty_pdf(self, pages: int = 1) -> Path:
        from pypdf import PdfWriter

        writer = PdfWriter()
        for _ in range(pages):
            writer.add_blank_page(width=200, height=200)
        path = self.root / "escaneado.pdf"
        with path.open("wb") as output:
            writer.write(output)
        return path

    def test_pdf_sem_camada_de_texto_sem_ocr_pede_outra_via(self):
        path = self._empty_pdf()
        with patch("audiofy.file_extraction.ocr_available", return_value=False):
            result = extract_file(path)
        self.assertTrue(result.needs_fallback)
        self.assertEqual(result.text, "")
        self.assertIn("escaneado", result.reason)
        self.assertIn("Tesseract", result.reason)

    def test_pdf_escaneado_tenta_ocr_quando_disponivel(self):
        path = self._empty_pdf()
        ocr = ExtractionResult("t", "texto reconhecido pelo ocr", "tesseract-ocr")
        with (
            patch("audiofy.file_extraction.ocr_available", return_value=True),
            patch("audiofy.file_extraction._ocr_pdf", return_value=ocr) as ocr_pdf,
        ):
            result = extract_file(path)
        ocr_pdf.assert_called_once()
        self.assertEqual(result.method, "tesseract-ocr")


class ImagemTest(unittest.TestCase):
    def setUp(self):
        self.directory = tempfile.TemporaryDirectory()
        self.addCleanup(self.directory.cleanup)
        self.path = Path(self.directory.name) / "pagina.png"
        from PIL import Image

        Image.new("RGB", (60, 30), "white").save(self.path)

    def test_imagem_sem_ocr_local_pede_outra_via_citando_setup(self):
        with patch("audiofy.file_extraction.ocr_available", return_value=False):
            result = extract_file(self.path)
        self.assertTrue(result.needs_fallback)
        self.assertEqual(result.text, "")
        self.assertIn("Tesseract", result.reason)
        self.assertIn("Configurações", result.reason)

    def test_ocr_bem_sucedido_retorna_o_texto_reconhecido(self):
        import pytesseract

        with (
            patch("audiofy.file_extraction.ocr_available", return_value=True),
            patch("audiofy.file_extraction._ocr_languages", return_value="por"),
            patch.object(pytesseract, "image_to_string", return_value=LONGO),
        ):
            result = _extract_image(self.path)
        self.assertEqual(result.method, "tesseract-ocr")
        self.assertFalse(result.needs_fallback)
        self.assertIn("Primeiro parágrafo", result.text)

    def test_ocr_com_pouco_texto_ainda_pede_outra_via(self):
        import pytesseract

        with (
            patch("audiofy.file_extraction.ocr_available", return_value=True),
            patch("audiofy.file_extraction._ocr_languages", return_value="eng"),
            patch.object(pytesseract, "image_to_string", return_value="oi"),
        ):
            result = _extract_image(self.path)
        self.assertTrue(result.needs_fallback)
        self.assertIn("não reconheceu texto suficiente", result.reason)

    def test_imagem_ilegivel_vira_erro_legivel(self):
        quebrada = Path(self.directory.name) / "quebrada.png"
        quebrada.write_bytes(b"nao e uma imagem de verdade")
        with patch("audiofy.file_extraction.ocr_available", return_value=True):
            with self.assertRaisesRegex(ValueError, "processar a imagem"):
                extract_file(quebrada)


class OcrPdfTest(unittest.TestCase):
    class _FakeImage:
        def __init__(self, data: bytes):
            self.data = data

    class _FakePage:
        def __init__(self, images):
            self.images = images

    class _FakeReader:
        def __init__(self, pages):
            self.pages = pages

    def _png_bytes(self) -> bytes:
        import io

        from PIL import Image

        buffer = io.BytesIO()
        Image.new("RGB", (40, 20), "white").save(buffer, format="PNG")
        return buffer.getvalue()

    def test_concatena_o_texto_reconhecido_de_cada_pagina(self):
        import pytesseract

        png = self._png_bytes()
        reader = self._FakeReader(
            [self._FakePage([self._FakeImage(png)]), self._FakePage([self._FakeImage(png)])]
        )
        with (
            patch("audiofy.file_extraction._ocr_languages", return_value="por+eng"),
            patch.object(pytesseract, "image_to_string", return_value=LONGO),
        ):
            result = _ocr_pdf(Path("x.pdf"), "Doc", reader)
        self.assertEqual(result.method, "tesseract-ocr")
        self.assertFalse(result.needs_fallback)
        self.assertEqual(result.text.count("Primeiro parágrafo"), 2)

    def test_pagina_com_imagem_ilegivel_nao_derruba_as_demais(self):
        import pytesseract

        reader = self._FakeReader(
            [
                self._FakePage([self._FakeImage(b"lixo")]),
                self._FakePage([self._FakeImage(self._png_bytes())]),
            ]
        )
        with (
            patch("audiofy.file_extraction._ocr_languages", return_value="por"),
            patch.object(pytesseract, "image_to_string", return_value=LONGO),
        ):
            result = _ocr_pdf(Path("x.pdf"), "Doc", reader)
        self.assertFalse(result.needs_fallback)
        self.assertEqual(result.text.count("Primeiro parágrafo"), 1)

    def test_ocr_vazio_no_pdf_pede_outra_via(self):
        reader = self._FakeReader([self._FakePage([])])
        with patch("audiofy.file_extraction._ocr_languages", return_value="eng"):
            result = _ocr_pdf(Path("x.pdf"), "Doc", reader)
        self.assertTrue(result.needs_fallback)
        self.assertIn("não reconheceu texto suficiente", result.reason)


class OcrIdiomaTest(unittest.TestCase):
    def test_prefere_portugues_e_ingles_quando_ambos_existem(self):
        import pytesseract

        with patch.object(pytesseract, "get_languages", return_value=["eng", "por", "deu"]):
            self.assertEqual(_ocr_languages(), "por+eng")

    def test_cai_para_ingles_quando_o_portugues_nao_esta_instalado(self):
        import pytesseract

        with patch.object(pytesseract, "get_languages", return_value=["eng", "deu"]):
            self.assertEqual(_ocr_languages(), "eng")

    def test_falha_ao_listar_idiomas_nao_interrompe_o_ocr(self):
        import pytesseract

        with patch.object(pytesseract, "get_languages", side_effect=OSError("tesseract sumiu")):
            self.assertEqual(_ocr_languages(), "eng")


class OcrDisponibilidadeTest(unittest.TestCase):
    def test_sem_binario_tesseract_o_ocr_fica_indisponivel(self):
        with patch("audiofy.setup.configure_tesseract", return_value=None):
            self.assertFalse(ocr_available())

    def test_binario_presente_e_pytesseract_instalado_habilita_ocr(self):
        with patch("audiofy.setup.configure_tesseract", return_value="/usr/bin/tesseract"):
            self.assertTrue(ocr_available())


if __name__ == "__main__":
    unittest.main()
