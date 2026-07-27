"""Extração de texto de arquivos locais por código, sem custo de IA.

Ordem de prioridade deliberada: Poppler (`pdftotext`) para PDF e bibliotecas
Python puras para os demais formatos (python-docx, ebooklib), OCR local
(Tesseract) para imagens e PDFs escaneados, e somente quando nada disso
funciona o chamador decide se aciona um modelo de IA — arquivos grandes
(livros, centenas de páginas) tornariam a via de IA lenta e cara, então ela
nunca é automática.

Em PDF o Poppler vem antes do `pypdf` por qualidade de extração: o `pypdf`
parte palavras em livros diagramados ("cav a-\\nleiro" em vez de "cavaleiro"),
defeito que atravessa o pipeline e chega no áudio. Como a extração é feita na
camada de texto do PDF, o resultado independe do idioma do documento.
"""

from __future__ import annotations

import html
import re
import shutil
import subprocess
from dataclasses import dataclass
from pathlib import Path

MAX_FILE_BYTES = 200 * 1024 * 1024
# Livro de centenas de páginas extrai em segundos; o teto só evita travar a UI
# se o binário ficar preso em um arquivo malformado.
_PDFTOTEXT_TIMEOUT_SECONDS = 180
_TEXT_SUFFIXES = {".txt", ".md", ".markdown", ".text"}
_IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff", ".bmp"}
SUPPORTED_SUFFIXES = {".pdf", ".docx", ".epub"} | _TEXT_SUFFIXES | _IMAGE_SUFFIXES

# Menos que isso por página indica PDF escaneado (só imagens, sem camada de texto).
_MIN_CHARS_PER_PAGE = 20
_MIN_TOTAL_CHARS = 120


@dataclass(frozen=True)
class ExtractionResult:
    """Resultado da extração: texto pronto ou o motivo de precisar de outra via."""

    title: str
    text: str
    method: str  # "pypdf" | "python-docx" | "ebooklib" | "plain-text" | "tesseract-ocr"
    # Quando needs_fallback=True, text está vazio e reason explica o que faltou.
    needs_fallback: bool = False
    reason: str = ""


def ocr_available() -> bool:
    """OCR local exige o binário Tesseract e a ponte pytesseract."""
    from .setup import configure_tesseract

    if not configure_tesseract():
        return False
    try:
        import pytesseract  # noqa: F401
        from PIL import Image  # noqa: F401
    except ImportError:
        return False
    return True


def _ocr_languages() -> str:
    """Usa português quando o traineddata existe; caso contrário o padrão do Tesseract."""
    import pytesseract

    try:
        available = set(pytesseract.get_languages(config=""))
    except Exception:
        return "eng"
    wanted = [lang for lang in ("por", "eng") if lang in available]
    return "+".join(wanted) if wanted else "eng"


def _normalize(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+\n", "\n", text)
    # Rodapé de diagramação InDesign ("...arquivo.indd   11 15/02/21   15:07").
    text = re.sub(
        r"^.*\.indd\s+\d+\s+\d{2}/\d{2}/\d{2}\s+\d{2}:\d{2}\s*$",
        "",
        text,
        flags=re.MULTILINE,
    )
    # Palavras partidas entre páginas pelo hifenizador (‑ U+2011 + quebra de linha).
    text = re.sub(r"\u2011\n", "", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Números de página soltos entre parágrafos ("...texto\n\n42\nMais texto...").
    text = re.sub(r"(?<=\n\n)\d{1,4}\n", "", text)
    return text.strip()


def _title_from_path(path: Path) -> str:
    return re.sub(r"[-_]+", " ", path.stem).strip() or path.stem


def extract_file(path: Path) -> ExtractionResult:
    """Extrai texto do arquivo pela melhor via local disponível."""
    if not path.is_file():
        raise FileNotFoundError(f"Arquivo não encontrado: {path}")
    size = path.stat().st_size
    if size == 0:
        raise ValueError("O arquivo está vazio.")
    if size > MAX_FILE_BYTES:
        raise ValueError("O arquivo excede o limite de 200 MiB.")
    suffix = path.suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return _extract_plain_text(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".epub":
        return _extract_epub(path)
    if suffix in _IMAGE_SUFFIXES:
        return _extract_image(path)
    raise ValueError(
        f"Formato '{suffix}' não suportado. Aceitos: PDF, DOCX, EPUB, TXT/MD e imagens."
    )


def _extract_plain_text(path: Path) -> ExtractionResult:
    raw = path.read_bytes()
    # UTF-16 só entra na disputa com BOM: sem ele, qualquer texto de bytes pares
    # "decodifica" em ideogramas silenciosamente e venceria o latin-1 legítimo.
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        candidates = ("utf-16", "utf-8", "latin-1")
    else:
        candidates = ("utf-8", "latin-1")
    for encoding in candidates:
        try:
            text = raw.decode(encoding)
            break
        except (UnicodeDecodeError, UnicodeError):
            continue
    else:  # latin-1 nunca falha; inalcançável, mas defensivo.
        text = raw.decode("utf-8", errors="replace")
    text = _normalize(text)
    if not text:
        raise ValueError("O arquivo de texto está vazio.")
    return ExtractionResult(_title_from_path(path), text, "plain-text")


_PAGE_NUMBER_RE = re.compile(r"\d+")
# Rodapé de diagramação é uma linha curta; frases de texto corrido são mais
# longas. O teto mantém fora da detecção uma abertura padronizada de capítulo,
# que se repetiria entre páginas mas é conteúdo que o ouvinte quer escutar.
_MAX_RUNNING_LINE_CHARS = 90
# Linha terminada em ponto final é frase, não marca de página.
_SENTENCE_END_RE = re.compile(r"[.!?][\"'”’)\]]*$")


def _repeated_line_signature(line: str) -> str:
    """Assinatura da linha ignorando números, para casar rodapés que só mudam a página.

    Retorna vazio quando a linha não tem cara de cabeçalho/rodapé, o que a mantém
    fora da detecção por repetição.
    """
    collapsed = " ".join(line.split())
    if (
        not collapsed
        or len(collapsed) > _MAX_RUNNING_LINE_CHARS
        or _SENTENCE_END_RE.search(collapsed)
    ):
        return ""
    return _PAGE_NUMBER_RE.sub("#", collapsed)


def _strip_running_headers(pages: list[str]) -> list[str]:
    """Remove cabeçalhos e rodapés que se repetem ao longo do documento.

    Diagramadores repetem título, nome do arquivo e numeração em toda página.
    Esse texto não é conteúdo: lido em voz alta vira ruído e, quando um trecho
    fica só com ele, o TTS chega a não devolver áudio nenhum. A detecção é por
    repetição — as bordas de cada página que aparecem em muitas páginas saem.
    """
    if len(pages) < 4:
        return pages

    def edge_positions(total: int) -> set[int]:
        """Bordas da página: até 2 linhas de cada ponta, nunca a página toda.

        Se as duas bordas se encostassem, o miolo — que é conteúdo — entraria na
        conta e poderia ser removido junto com o cabeçalho.
        """
        if total <= 1:
            return {0} if total == 1 else set()
        margin = 1 if total < 6 else 2
        return set(range(margin)) | set(range(total - margin, total))

    edge_counts: dict[str, int] = {}
    for page in pages:
        lines = [line for line in page.splitlines() if line.strip()]
        for position in edge_positions(len(lines)):
            signature = _repeated_line_signature(lines[position])
            if signature:
                edge_counts[signature] = edge_counts.get(signature, 0) + 1
    threshold = max(3, len(pages) // 3)
    running = {signature for signature, count in edge_counts.items() if count >= threshold}
    if not running:
        return pages
    cleaned: list[str] = []
    for page in pages:
        lines = page.splitlines()
        borders = edge_positions(len(lines))
        keep = [
            line
            for position, line in enumerate(lines)
            if not (position in borders and _repeated_line_signature(line) in running)
        ]
        cleaned.append("\n".join(keep))
    return cleaned


def _pdftotext_binary() -> str | None:
    """Caminho do `pdftotext` (pacote poppler-utils), ou None se não instalado."""
    return shutil.which("pdftotext")


def _run_pdftotext(binary: str, path: Path) -> str:
    """Extrai o texto do PDF pelo Poppler, preservando a ordem de leitura.

    `-layout` manteria o alinhamento visual em colunas, o que atrapalha leitura
    em voz alta; o padrão (fluxo contínuo) é o que queremos para TTS.
    """
    completed = subprocess.run(
        [binary, "-enc", "UTF-8", "-nopgbrk", str(path), "-"],
        capture_output=True,
        timeout=_PDFTOTEXT_TIMEOUT_SECONDS,
        check=True,
    )
    return completed.stdout.decode("utf-8", errors="replace")


def _extract_pdf_with_pypdf(path: Path) -> str:
    """Via reserva: só entra quando o Poppler não está disponível ou falha."""
    reader = _open_pdf(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return _normalize("\n\n".join(_strip_running_headers(pages)))


def _extract_pdf_text(path: Path) -> tuple[str, str]:
    """Devolve (método, texto) da melhor via disponível para PDF com texto.

    O Poppler vem primeiro porque o `pypdf` parte palavras em PDFs de livro
    diagramado: no PDF real de "O cavaleiro preso na armadura" ele produzia 276
    quebras ("cav a-\\nleiro", "men ção", "Chris topher"), que atravessavam o
    pipeline inteiro e chegavam quebradas no áudio. O mesmo arquivo pelo
    pdftotext dá zero — a correção acontece na origem, sem heurística depois.
    Por operar na camada de texto do PDF, vale para qualquer idioma.
    """
    binary = _pdftotext_binary()
    if binary:
        try:
            text = _normalize(_run_pdftotext(binary, path))
            if text.strip():
                return "pdftotext", text
        except (OSError, subprocess.SubprocessError):
            pass  # Poppler indisponível/falhou: a via reserva assume.
    return "pypdf", _extract_pdf_with_pypdf(path)


def _open_pdf(path: Path):
    try:
        from pypdf import PdfReader
    except ImportError as error:
        raise RuntimeError(
            "A biblioteca pypdf não está instalada; rode Instalar/corrigir em Configurações."
        ) from error
    try:
        reader = PdfReader(path)
    except Exception as error:
        raise ValueError(f"Não consegui abrir o PDF: {error}") from error
    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception as error:
            raise ValueError("O PDF é protegido por senha.") from error
    return reader


def _extract_pdf(path: Path) -> ExtractionResult:
    reader = _open_pdf(path)
    method, text = _extract_pdf_text(path)
    title = ""
    if reader.metadata and reader.metadata.title:
        title = str(reader.metadata.title).strip()
    title = title or _title_from_path(path)
    # Camada de texto ausente/ínfima = PDF escaneado; texto "extraído" seria lixo.
    if len(text) < max(_MIN_TOTAL_CHARS, _MIN_CHARS_PER_PAGE * len(reader.pages)):
        if ocr_available():
            return _ocr_pdf(path, title, reader)
        return ExtractionResult(
            title,
            "",
            method,
            needs_fallback=True,
            reason=(
                "O PDF parece ser escaneado (sem camada de texto) e o OCR local "
                "(Tesseract) não está instalado."
            ),
        )
    return ExtractionResult(title, text, method)


def _ocr_pdf(path: Path, title: str, reader) -> ExtractionResult:
    """OCR página a página de um PDF escaneado usando as imagens embutidas."""
    import io

    import pytesseract
    from PIL import Image

    languages = _ocr_languages()
    texts: list[str] = []
    for page in reader.pages:
        for image_file in page.images:
            try:
                with Image.open(io.BytesIO(image_file.data)) as image:
                    texts.append(pytesseract.image_to_string(image, lang=languages))
            except Exception:  # página sem imagem legível não derruba o restante
                continue
    text = _normalize("\n\n".join(texts))
    if len(text) < _MIN_TOTAL_CHARS:
        return ExtractionResult(
            title,
            "",
            "tesseract-ocr",
            needs_fallback=True,
            reason="O OCR local não reconheceu texto suficiente nas páginas do PDF.",
        )
    return ExtractionResult(title, text, "tesseract-ocr")


def _extract_docx(path: Path) -> ExtractionResult:
    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError(
            "A biblioteca python-docx não está instalada; rode Instalar/corrigir em Configurações."
        ) from error
    try:
        document = Document(str(path))
    except Exception as error:
        raise ValueError(f"Não consegui abrir o DOCX: {error}") from error
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" — ".join(cells))
    text = _normalize("\n\n".join(part for part in parts if part.strip()))
    if len(text) < _MIN_TOTAL_CHARS:
        return ExtractionResult(
            _title_from_path(path),
            "",
            "python-docx",
            needs_fallback=True,
            reason="O DOCX não contém texto extraível (pode ser só imagens).",
        )
    title = (document.core_properties.title or "").strip() or _title_from_path(path)
    return ExtractionResult(title, text, "python-docx")


_EPUB_TAG_RE = re.compile(r"<[^>]+>")
_EPUB_BLOCK_RE = re.compile(
    r"</(?:p|div|h[1-6]|li|blockquote|section|article)>|<br\s*/?>", re.IGNORECASE
)


def _epub_html_to_text(markup: str) -> str:
    markup = re.sub(
        r"<(?:script|style)[^>]*>.*?</(?:script|style)>",
        "",
        markup,
        flags=re.DOTALL | re.IGNORECASE,
    )
    markup = _EPUB_BLOCK_RE.sub("\n\n", markup)
    return html.unescape(_EPUB_TAG_RE.sub("", markup))


def _extract_epub(path: Path) -> ExtractionResult:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except ImportError as error:
        raise RuntimeError(
            "A biblioteca ebooklib não está instalada; rode Instalar/corrigir em Configurações."
        ) from error
    try:
        book = epub.read_epub(str(path), options={"ignore_ncx": True})
    except Exception as error:
        raise ValueError(f"Não consegui abrir o EPUB: {error}") from error
    chapters = [
        _epub_html_to_text(item.get_content().decode("utf-8", errors="replace"))
        for item in book.get_items_of_type(ITEM_DOCUMENT)
    ]
    text = _normalize("\n\n".join(chapters))
    if len(text) < _MIN_TOTAL_CHARS:
        return ExtractionResult(
            _title_from_path(path),
            "",
            "ebooklib",
            needs_fallback=True,
            reason="O EPUB não contém texto extraível.",
        )
    title = (book.get_metadata("DC", "title") or [("", {})])[0][0] or _title_from_path(path)
    return ExtractionResult(str(title).strip() or _title_from_path(path), text, "ebooklib")


def _extract_image(path: Path) -> ExtractionResult:
    title = _title_from_path(path)
    if not ocr_available():
        return ExtractionResult(
            title,
            "",
            "tesseract-ocr",
            needs_fallback=True,
            reason=(
                "Ler texto de imagem exige o OCR local (Tesseract), que não está "
                "instalado. Veja Diagnóstico/Setup em Configurações."
            ),
        )
    import pytesseract
    from PIL import Image

    try:
        with Image.open(path) as image:
            text = _normalize(pytesseract.image_to_string(image, lang=_ocr_languages()))
    except Exception as error:
        raise ValueError(f"Não consegui processar a imagem: {error}") from error
    if len(text) < _MIN_TOTAL_CHARS:
        return ExtractionResult(
            title,
            "",
            "tesseract-ocr",
            needs_fallback=True,
            reason="O OCR local não reconheceu texto suficiente na imagem.",
        )
    return ExtractionResult(title, text, "tesseract-ocr")
